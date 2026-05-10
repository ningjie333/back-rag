"""多格式文件解析 + 文本分块

支持格式：PDF, TXT, MD, DOCX, DOC, XLSX, XLS, CSV, JSON, YAML
参考 graph-rag-agent-0521-master/processor/file_reader.py 的架构设计
"""
import codecs
import csv
import json
import os
import re
from pathlib import Path
from typing import Iterator

import pymupdf
import yaml
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config import settings

SUPPORTED_EXTS = {".pdf", ".txt", ".md", ".docx", ".doc", ".xlsx", ".xls", ".csv", ".json", ".yaml", ".yml"}

CHUNK_SEPARATORS = ["\n\n", "\n", "。", "！", "？", " ", ""]


def get_text_splitter() -> RecursiveCharacterTextSplitter:
    return RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        separators=CHUNK_SEPARATORS,
    )


def chunk_text(text: str) -> list[str]:
    splitter = get_text_splitter()
    return splitter.split_text(text)


def _read_txt_with_fallback(file_path: str) -> str:
    """读取文本文件，带编码自动检测"""
    try:
        with codecs.open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    except Exception:
        pass
    try:
        with open(file_path, "rb") as f:
            raw = f.read(10240)
        try:
            import chardet
            enc = chardet.detect(raw).get("encoding") or "gbk"
        except ImportError:
            enc = "gbk"
        with codecs.open(file_path, "r", encoding=enc, errors="replace") as f:
            return f.read()
    except Exception as e:
        return f"[无法读取文件内容: {e}]"


def _read_pdf(file_path: str) -> Iterator[dict]:
    """逐页提取 PDF，带单页错误恢复"""
    try:
        doc = pymupdf.open(str(file_path))
        for i in range(len(doc)):
            try:
                text = doc[i].get_text()
                if text.strip():
                    yield {"page": i + 1, "text": text.strip()}
            except Exception:
                yield {"page": i + 1, "text": f"[第 {i+1} 页无法读取]"}
        doc.close()
    except Exception as e:
        yield {"page": 0, "text": f"[无法读取PDF: {e}]"}


def _read_docx(file_path: str) -> str:
    """读取 DOCX"""
    try:
        from docx import Document
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        return f"[无法读取DOCX: {e}]"


def _read_doc(file_path: str) -> str:
    """读取旧版 DOC，三级降级策略"""
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(file_path))
        content = doc.Content.Text
        doc.Close()
        word.Quit()
        if content.strip():
            return content
    except (ImportError, Exception):
        pass
    try:
        import textract
        content = textract.process(file_path).decode("utf-8")
        if content.strip():
            return content
    except (ImportError, Exception):
        pass
    try:
        from docx import Document
        doc = Document(file_path)
        content = "\n".join(p.text for p in doc.paragraphs)
        if content.strip():
            return content
    except (ImportError, Exception):
        pass
    return "[无法读取DOC文件，建议转换为DOCX格式]"


def _read_excel(file_path: str) -> str:
    """读取 Excel (.xlsx/.xls)，按 sheet 拼接为文本"""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"## Sheet: {sheet_name}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(cell) if cell is not None else "" for cell in row]
                parts.append("\t".join(cells))
            parts.append("")
        wb.close()
        return "\n".join(parts)
    except ImportError:
        return "[无法读取Excel: 缺少 openpyxl 库，请安装 pip install openpyxl]"
    except Exception as e:
        return f"[无法读取Excel: {e}]"


def _read_csv_with_fallback(file_path: str) -> str:
    """读取 CSV，带编码回退"""
    for enc in ("utf-8", "gbk"):
        try:
            rows = []
            with open(file_path, "r", encoding=enc, errors="replace") as f:
                reader = csv.reader(f)
                for row in reader:
                    rows.append(",".join(row))
            return "\n".join(rows)
        except Exception:
            continue
    return "[无法读取CSV文件]"


def _read_json(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            data = json.load(f)
        return json.dumps(data, ensure_ascii=False, indent=2)
    except Exception as e:
        return f"[无法读取JSON: {e}]"


def _read_yaml(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            data = yaml.safe_load(f)
        return yaml.dump(data, allow_unicode=True, default_flow_style=False)
    except Exception as e:
        return f"[无法读取YAML: {e}]"


def extract_text(file_path: str | Path, ext: str | None = None) -> Iterator[dict]:
    """
    通用文本提取入口，逐页/逐段返回。

    Returns:
        Iterator[dict]: {page: int, text: str}
    """
    path = Path(file_path)
    ext = ext or path.suffix.lower()

    if ext == ".pdf":
        yield from _read_pdf(str(path))
    else:
        text = ""
        if ext in (".txt", ".md"):
            text = _read_txt_with_fallback(str(path))
        elif ext == ".docx":
            text = _read_docx(str(path))
        elif ext == ".doc":
            text = _read_doc(str(path))
        elif ext in (".xlsx", ".xls"):
            text = _read_excel(str(path))
        elif ext == ".csv":
            text = _read_csv_with_fallback(str(path))
        elif ext == ".json":
            text = _read_json(str(path))
        elif ext in (".yaml", ".yml"):
            text = _read_yaml(str(path))
        else:
            text = f"[不支持的文件格式: {ext}]"

        if text.strip():
            yield {"page": 1, "text": text.strip()}


def process_file_to_chunks(
    file_path: str | Path,
    book_title: str,
    ext: str | None = None,
) -> list[dict]:
    """
    处理任意支持的文件 → 分块列表。

    Returns:
        list[dict]: [{chunk_id, text, source, page, chunk_index, book_title}]
    """
    path = Path(file_path)
    ext = ext or path.suffix.lower()
    stem = path.stem

    chunks: list[dict] = []
    chunk_index = 0

    for page_data in extract_text(path, ext):
        pieces = chunk_text(page_data["text"])
        for piece in pieces:
            if len(piece.strip()) < 20:
                continue
            chunks.append({
                "chunk_id": f"{stem}_p{page_data['page']}_c{chunk_index}",
                "text": piece,
                "source": str(path),
                "page": page_data["page"],
                "chapter": "",
                "book_title": book_title,
                "chunk_index": chunk_index,
            })
            chunk_index += 1

    return chunks
